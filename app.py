import io
import json
import os
import uuid
import hashlib
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, TypedDict

import streamlit as st
from dotenv import load_dotenv
from docx import Document as DocxDocument
from pypdf import PdfReader

from sqlalchemy import Boolean, DateTime, ForeignKey, String, Text, create_engine, select
from sqlalchemy.orm import DeclarativeBase, Mapped, Session, mapped_column, relationship

from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langgraph.graph import END, START, StateGraph


# ------------------------------
# Configuration
# ------------------------------
load_dotenv(dotenv_path=Path(".env"), override=False)

DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./app.db")
CHROMA_BASE_DIR = "./chroma_db"

DEFAULT_SYSTEM_PROMPT = (
    "You are a helpful assistant. If RAG context is provided, prioritize it. "
    "If sufficient context is missing, say you don't know. "
    "Do not guess or fabricate facts. "
    "Keep answers concise and accurate."
)


# ------------------------------
# Database models
# ------------------------------
class Base(DeclarativeBase):
    pass


class Chat(Base):
    __tablename__ = "chats"

    id: Mapped[str] = mapped_column(String(36), primary_key=True)
    title: Mapped[str] = mapped_column(String(200), nullable=False)
    system_prompt: Mapped[str] = mapped_column(Text, nullable=False, default=DEFAULT_SYSTEM_PROMPT)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), nullable=False)
    updated_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), nullable=False)

    messages: Mapped[List["Message"]] = relationship(back_populates="chat", cascade="all, delete-orphan")
    documents: Mapped[List["DocumentRecord"]] = relationship(back_populates="chat", cascade="all, delete-orphan")


class Message(Base):
    __tablename__ = "messages"

    id: Mapped[str] = mapped_column(String(36), primary_key=True)
    chat_id: Mapped[str] = mapped_column(ForeignKey("chats.id"), index=True)
    role: Mapped[str] = mapped_column(String(20), nullable=False)
    content: Mapped[str] = mapped_column(Text, nullable=False)
    used_rag: Mapped[bool] = mapped_column(Boolean, nullable=False, default=False)
    evidence_json: Mapped[str] = mapped_column(Text, nullable=False, default="[]")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), nullable=False)

    chat: Mapped[Chat] = relationship(back_populates="messages")


class DocumentRecord(Base):
    __tablename__ = "documents"

    id: Mapped[str] = mapped_column(String(36), primary_key=True)
    chat_id: Mapped[str] = mapped_column(ForeignKey("chats.id"), index=True)
    doc_id: Mapped[str] = mapped_column(String(36), index=True)
    filename: Mapped[str] = mapped_column(String(500), nullable=False)
    content_hash: Mapped[str] = mapped_column(String(64), nullable=False, index=True)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), nullable=False)

    chat: Mapped[Chat] = relationship(back_populates="documents")


engine = create_engine(DATABASE_URL, future=True)
Base.metadata.create_all(engine)


def now_utc() -> datetime:
    return datetime.now(timezone.utc)


def get_db() -> Session:
    return Session(engine)


# ------------------------------
# Chat persistence helpers
# ------------------------------
def create_chat(title: str) -> str:
    chat_id = str(uuid.uuid4())
    ts = now_utc()
    with get_db() as db:
        db.add(
            Chat(
                id=chat_id,
                title=title,
                system_prompt=DEFAULT_SYSTEM_PROMPT,
                created_at=ts,
                updated_at=ts,
            )
        )
        db.commit()
    return chat_id


def get_or_create_default_chat() -> str:
    with get_db() as db:
        first_chat = db.scalar(select(Chat).order_by(Chat.created_at.asc()))
        if first_chat:
            return first_chat.id
    return create_chat("Chat 1")


def list_chats() -> List[Chat]:
    with get_db() as db:
        return list(db.scalars(select(Chat).order_by(Chat.updated_at.desc())))


def get_chat(chat_id: str) -> Chat | None:
    with get_db() as db:
        return db.get(Chat, chat_id)


def update_chat_prompt(chat_id: str, prompt: str):
    with get_db() as db:
        chat = db.get(Chat, chat_id)
        if not chat:
            return
        chat.system_prompt = prompt
        chat.updated_at = now_utc()
        db.commit()


def update_chat_title(chat_id: str, title: str):
    with get_db() as db:
        chat = db.get(Chat, chat_id)
        if not chat:
            return
        chat.title = title
        chat.updated_at = now_utc()
        db.commit()


def delete_chat(chat_id: str):
    with get_db() as db:
        chat = db.get(Chat, chat_id)
        if not chat:
            return
        db.delete(chat)
        db.commit()
    chat_chroma_dir = Path(CHROMA_BASE_DIR) / chat_id
    shutil.rmtree(chat_chroma_dir, ignore_errors=True)


def append_message(chat_id: str, role: str, content: str, used_rag: bool = False, evidence: List[Dict[str, Any]] | None = None):
    with get_db() as db:
        chat = db.get(Chat, chat_id)
        if not chat:
            return
        db.add(
            Message(
                id=str(uuid.uuid4()),
                chat_id=chat_id,
                role=role,
                content=content,
                used_rag=used_rag,
                evidence_json=json.dumps(evidence or []),
                created_at=now_utc(),
            )
        )
        chat.updated_at = now_utc()
        db.commit()


def clear_messages(chat_id: str):
    with get_db() as db:
        chat = db.get(Chat, chat_id)
        if not chat:
            return
        for msg in list(chat.messages):
            db.delete(msg)
        chat.updated_at = now_utc()
        db.commit()


def get_messages(chat_id: str) -> List[Dict[str, Any]]:
    with get_db() as db:
        rows = list(db.scalars(select(Message).where(Message.chat_id == chat_id).order_by(Message.created_at.asc())))
    out: List[Dict[str, Any]] = []
    for row in rows:
        out.append(
            {
                "role": row.role,
                "content": row.content,
                "used_rag": row.used_rag,
                "evidence": json.loads(row.evidence_json or "[]"),
            }
        )
    return out


def is_duplicate_hash(chat_id: str, content_hash: str) -> bool:
    with get_db() as db:
        existing = db.scalar(
            select(DocumentRecord).where(DocumentRecord.chat_id == chat_id, DocumentRecord.content_hash == content_hash)
        )
        return existing is not None


def add_document_record(chat_id: str, doc_id: str, filename: str, content_hash: str):
    with get_db() as db:
        chat = db.get(Chat, chat_id)
        if not chat:
            return
        db.add(
            DocumentRecord(
                id=str(uuid.uuid4()),
                chat_id=chat_id,
                doc_id=doc_id,
                filename=filename,
                content_hash=content_hash,
                created_at=now_utc(),
            )
        )
        chat.updated_at = now_utc()
        db.commit()


def list_documents(chat_id: str) -> Dict[str, str]:
    with get_db() as db:
        rows = list(
            db.scalars(select(DocumentRecord).where(DocumentRecord.chat_id == chat_id).order_by(DocumentRecord.created_at.asc()))
        )
    docs: Dict[str, str] = {}
    for row in rows:
        if row.doc_id not in docs:
            docs[row.doc_id] = row.filename
    return docs


def delete_document_records(chat_id: str, doc_id: str):
    with get_db() as db:
        chat = db.get(Chat, chat_id)
        if not chat:
            return
        rows = list(
            db.scalars(select(DocumentRecord).where(DocumentRecord.chat_id == chat_id, DocumentRecord.doc_id == doc_id))
        )
        for row in rows:
            db.delete(row)
        chat.updated_at = now_utc()
        db.commit()


# ------------------------------
# Helpers: parsing and ingestion
# ------------------------------
def parse_file(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix in {".txt", ".md", ".csv"}:
        return file_bytes.decode("utf-8", errors="ignore")

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages)

    if suffix == ".docx":
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)

    raise ValueError(f"Unsupported file type: {suffix}")


def get_vectorstore(chat_id: str) -> Chroma:
    chat_chroma_dir = Path(CHROMA_BASE_DIR) / chat_id
    chat_chroma_dir.mkdir(parents=True, exist_ok=True)
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    return Chroma(
        collection_name="uploaded_docs",
        embedding_function=embeddings,
        persist_directory=str(chat_chroma_dir),
    )


def file_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def ingest_documents(chat_id: str, vectorstore: Chroma, uploaded_files) -> Dict[str, int]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    ingested_count = 0
    skipped_duplicates = 0

    for file in uploaded_files:
        file_bytes = file.getvalue()
        content_hash = file_hash(file_bytes)
        if is_duplicate_hash(chat_id, content_hash):
            skipped_duplicates += 1
            continue

        text = parse_file(file_bytes, file.name)
        if not text.strip():
            continue

        doc_id = str(uuid.uuid4())
        chunks = splitter.split_text(text)

        docs = []
        ids = []
        for i, chunk in enumerate(chunks):
            chunk_id = f"{doc_id}-{i}"
            docs.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "doc_id": doc_id,
                        "chat_id": chat_id,
                        "content_hash": content_hash,
                        "source": file.name,
                        "chunk_index": i,
                    },
                )
            )
            ids.append(chunk_id)

        vectorstore.add_documents(docs, ids=ids)
        add_document_record(chat_id, doc_id, file.name, content_hash)
        ingested_count += 1

    return {"ingested_count": ingested_count, "skipped_duplicates": skipped_duplicates}


def delete_document_from_vectorstore(vectorstore: Chroma, doc_id: str):
    vectorstore._collection.delete(where={"doc_id": doc_id})


def recent_chat_context(messages: List[Dict[str, Any]], limit: int = 6) -> str:
    lines = []
    for msg in messages[-limit:]:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        lines.append(f"{role.upper()}: {content}")
    return "\n".join(lines)


def docs_to_evidence(docs: List[Document]) -> List[Dict[str, Any]]:
    evidence = []
    for d in docs:
        evidence.append(
            {
                "source": d.metadata.get("source", "unknown"),
                "chunk_index": d.metadata.get("chunk_index", "?"),
                "content": d.page_content,
            }
        )
    return evidence


# ------------------------------
# LangGraph single-agent flow
# ------------------------------
class GraphState(TypedDict):
    question: str
    system_prompt: str
    chat_history: List[Dict[str, Any]]
    rag_mode: str
    docs_available: bool
    available_sources: List[str]
    use_rag: bool
    retrieved_docs: List[Document]
    answer: str


def build_graph(vectorstore: Chroma):
    retriever = vectorstore.as_retriever(search_kwargs={"k": 4})
    llm = ChatOpenAI(model="gpt-4.1-mini", temperature=0)

    def agent(state: GraphState):
        use_rag = False
        question = state["question"]
        question_lc = question.lower()
        rag_mode = state["rag_mode"]
        available_sources = state.get("available_sources", [])
        retrieved_docs: List[Document] = []

        file_list_intents = ["what documents", "which documents", "what files", "which files", "uploaded docs", "uploaded files"]
        if any(intent in question_lc for intent in file_list_intents):
            if available_sources:
                listed = "\n".join([f"- {name}" for name in available_sources])
                return {
                    "answer": f"I can see these uploaded documents:\n{listed}",
                    "retrieved_docs": [],
                    "use_rag": False,
                }
            return {
                "answer": "No uploaded documents are currently indexed.",
                "retrieved_docs": [],
                "use_rag": False,
            }

        if state["docs_available"]:
            if rag_mode == "always":
                use_rag = True
            elif rag_mode == "never":
                use_rag = False
            elif any(src.lower() in question_lc for src in available_sources if src):
                use_rag = True
            else:
                history = recent_chat_context(state.get("chat_history", []))
                sources_text = ", ".join(available_sources) if available_sources else "None"
                decision_prompt = (
                    "Decide whether retrieval from uploaded docs is necessary for the user's question. "
                    "Reply with only YES or NO.\n\n"
                    f"Uploaded file names:\n{sources_text}\n\n"
                    f"Recent chat:\n{history}\n\n"
                    f"Question:\n{state['question']}"
                )
                decision = llm.invoke(
                    [
                        {"role": "system", "content": "Return only YES or NO."},
                        {"role": "user", "content": decision_prompt},
                    ]
                )
                use_rag = str(decision.content).strip().upper().startswith("Y")

        if use_rag:
            retrieved_docs = retriever.invoke(state["question"])

        history = recent_chat_context(state.get("chat_history", []))
        if use_rag and retrieved_docs:
            context = "\n\n".join(
                [
                    (
                        f"[Source: {d.metadata.get('source', 'unknown')} | "
                        f"Chunk: {d.metadata.get('chunk_index', '?')}]\n{d.page_content}"
                    )
                    for d in retrieved_docs
                ]
            )
            user_prompt = (
                f"Recent chat:\n{history}\n\n"
                f"Question:\n{state['question']}\n\n"
                f"Context:\n{context}\n\n"
                "Use the context when answering. End with a short 'Sources used' list."
            )
        else:
            sources_text = ", ".join(available_sources) if available_sources else "None"
            user_prompt = (
                f"Recent chat:\n{history}\n\n"
                f"Uploaded file names:\n{sources_text}\n\n"
                f"Question:\n{question}\n\n"
                "Answer directly without using uploaded-document facts. "
                "If the question depends on uploaded docs or missing evidence, reply with: "
                "'I don't know based on the current context.' "
                "Do not guess or invent resume/profile/work-history details. "
                "Do not claim there are no uploaded docs if file names are provided above."
            )

        msg = llm.invoke(
            [
                {"role": "system", "content": state["system_prompt"]},
                {"role": "user", "content": user_prompt},
            ]
        )
        return {"answer": msg.content, "retrieved_docs": retrieved_docs, "use_rag": use_rag}

    graph = StateGraph(GraphState)
    graph.add_node("agent", agent)
    graph.add_edge(START, "agent")
    graph.add_edge("agent", END)
    return graph.compile()


# ------------------------------
# Streamlit UI
# ------------------------------
def chat_label(chat: Chat) -> str:
    ts = chat.updated_at.strftime("%Y-%m-%d %H:%M") if chat.updated_at else ""
    return f"{chat.title} ({ts})"


def main():
    st.set_page_config(page_title="RAG Multi-Chat Assistant", layout="wide")
    st.title("Document RAG Multi-Chat Assistant")

    if not os.getenv("OPENAI_API_KEY"):
        st.error(
            "Missing OPENAI_API_KEY. Create a `.env` file in the project root with:\n\n"
            "OPENAI_API_KEY=your_key_here"
        )
        st.stop()

    if "active_chat_id" not in st.session_state:
        st.session_state["active_chat_id"] = get_or_create_default_chat()
    if "uploader_versions" not in st.session_state:
        st.session_state["uploader_versions"] = {}

    chats = list_chats()
    if not chats:
        st.session_state["active_chat_id"] = create_chat("Chat 1")
        chats = list_chats()

    active_chat_id = st.session_state["active_chat_id"]
    if active_chat_id not in {c.id for c in chats}:
        active_chat_id = chats[0].id
        st.session_state["active_chat_id"] = active_chat_id

    chat = get_chat(active_chat_id)
    if not chat:
        st.error("Active chat not found.")
        st.stop()

    vectorstore = get_vectorstore(active_chat_id)
    graph = build_graph(vectorstore)

    @st.dialog("Create New Chat")
    def create_chat_dialog():
        default_name = f"Chat {len(chats) + 1}"
        new_name = st.text_input("Chat name", value=default_name, key="new_chat_name_input")
        col1, col2 = st.columns(2)
        if col1.button("Create", type="primary", use_container_width=True):
            final_name = new_name.strip() or default_name
            new_chat_id = create_chat(final_name)
            st.session_state["active_chat_id"] = new_chat_id
            st.rerun()
        if col2.button("Cancel", use_container_width=True):
            st.rerun()

    @st.dialog("Rename Chat")
    def rename_chat_dialog():
        renamed = st.text_input("New chat name", value=chat.title, key=f"rename_chat_name_{active_chat_id}")
        col1, col2 = st.columns(2)
        if col1.button("Save", type="primary", use_container_width=True):
            final_name = renamed.strip() or chat.title
            update_chat_title(active_chat_id, final_name)
            st.rerun()
        if col2.button("Cancel", use_container_width=True):
            st.rerun()

    @st.dialog("Delete Chat")
    def delete_chat_dialog():
        if len(chats) <= 1:
            st.warning("At least one chat must exist. Create another chat before deleting this one.")
            if st.button("Close", use_container_width=True):
                st.rerun()
            return

        st.write(f"Delete chat **{chat.title}**?")
        st.caption("This removes chat messages, document records, and that chat's vector DB.")
        col1, col2 = st.columns(2)
        if col1.button("Delete", type="primary", use_container_width=True):
            delete_chat(active_chat_id)
            remaining = list_chats()
            if remaining:
                st.session_state["active_chat_id"] = remaining[0].id
            else:
                st.session_state["active_chat_id"] = create_chat("Chat 1")
            st.session_state["uploader_versions"].pop(active_chat_id, None)
            st.rerun()
        if col2.button("Cancel", use_container_width=True):
            st.rerun()

    with st.sidebar:
        st.header("Chats")
        if st.button("New Chat", type="primary", use_container_width=True):
            create_chat_dialog()

        chat_options = {c.id: chat_label(c) for c in chats}
        selected_chat_id = st.selectbox(
            "Open Existing Chat",
            options=list(chat_options.keys()),
            index=list(chat_options.keys()).index(active_chat_id),
            format_func=lambda cid: chat_options[cid],
        )
        if selected_chat_id != active_chat_id:
            st.session_state["active_chat_id"] = selected_chat_id
            st.rerun()

        c1, c2 = st.columns(2)
        if c1.button("Rename", type="primary", use_container_width=True):
            rename_chat_dialog()
        if c2.button("Delete", type="primary", use_container_width=True):
            delete_chat_dialog()

        st.divider()
        st.header("Settings")

        prompt_value = st.text_area(
            "System Prompt",
            value=chat.system_prompt,
            height=180,
            help="This prompt is stored per chat.",
        )
        if prompt_value != chat.system_prompt:
            update_chat_prompt(active_chat_id, prompt_value)
            chat = get_chat(active_chat_id)

        rag_mode_ui = st.radio(
            "RAG Mode",
            options=["Auto decide", "Always refer documents", "Do not refer documents"],
            index=0,
        )
        rag_mode_map = {
            "Auto decide": "auto",
            "Always refer documents": "always",
            "Do not refer documents": "never",
        }
        rag_mode = rag_mode_map[rag_mode_ui]

        if st.button("Clear chat history", type="primary", use_container_width=True):
            clear_messages(active_chat_id)
            st.rerun()

        st.subheader("Upload Documents")
        upload_key = f"uploaded_files_{active_chat_id}_{st.session_state['uploader_versions'].get(active_chat_id, 0)}"
        uploaded_files = st.file_uploader(
            "Upload .txt, .md, .csv, .pdf, .docx",
            type=["txt", "md", "csv", "pdf", "docx"],
            accept_multiple_files=True,
            key=upload_key,
        )

        if st.button("Ingest Uploaded Files", type="primary", use_container_width=True):
            if not uploaded_files:
                st.warning("No files selected.")
            else:
                try:
                    stats = ingest_documents(active_chat_id, vectorstore, uploaded_files)
                    st.success(
                        f"Ingested {stats['ingested_count']} document(s). "
                        f"Skipped {stats['skipped_duplicates']} duplicate(s)."
                    )
                    st.session_state["uploader_versions"][active_chat_id] = (
                        st.session_state["uploader_versions"].get(active_chat_id, 0) + 1
                    )
                    st.rerun()
                except Exception as e:
                    st.exception(e)

        st.subheader("Manage Vector DB")
        documents = list_documents(active_chat_id)
        if not documents:
            st.caption("No documents in this chat's vector DB.")
        else:
            selected_doc_id = st.selectbox(
                "Select document to delete",
                options=list(documents.keys()),
                format_func=lambda x: f"{documents[x]} ({x[:8]}...)",
            )
            if st.button("Delete Selected Document", type="primary", use_container_width=True):
                delete_document_from_vectorstore(vectorstore, selected_doc_id)
                delete_document_records(active_chat_id, selected_doc_id)
                st.success("Document deleted from this chat.")
                st.rerun()

    messages = get_messages(active_chat_id)
    for msg in messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            evidence = msg.get("evidence", [])
            if msg["role"] == "assistant" and evidence:
                with st.expander("Retrieved Chunks (Evidence)"):
                    for idx, ev in enumerate(evidence, start=1):
                        st.markdown(f"**{idx}. {ev['source']} | chunk {ev['chunk_index']}**")
                        st.write(ev["content"])

    prompt = st.chat_input("Ask a question about your uploaded documents")
    if prompt:
        append_message(active_chat_id, role="user", content=prompt)
        with st.chat_message("user"):
            st.markdown(prompt)

        messages = get_messages(active_chat_id)
        chat_history = messages[:-1]

        docs = list_documents(active_chat_id)
        docs_available = len(docs) > 0
        available_sources = list(docs.values())

        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                result = graph.invoke(
                    {
                        "question": prompt,
                        "system_prompt": chat.system_prompt,
                        "chat_history": chat_history,
                        "rag_mode": rag_mode,
                        "docs_available": docs_available,
                        "available_sources": available_sources,
                        "use_rag": False,
                        "retrieved_docs": [],
                        "answer": "",
                    }
                )

            answer = result["answer"]
            evidence = docs_to_evidence(result.get("retrieved_docs", []))
            st.markdown(answer)
            if evidence:
                with st.expander("Retrieved Chunks (Evidence)"):
                    for idx, ev in enumerate(evidence, start=1):
                        st.markdown(f"**{idx}. {ev['source']} | chunk {ev['chunk_index']}**")
                        st.write(ev["content"])

        append_message(
            active_chat_id,
            role="assistant",
            content=answer,
            used_rag=bool(result.get("use_rag", False)),
            evidence=evidence,
        )


if __name__ == "__main__":
    main()
